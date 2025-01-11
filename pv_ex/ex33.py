# -*- coding: utf-8 -*-
# @Time : 2025/01/09 15:54
# @Author : zhanglei
# @Email : zhspark@gmail.com


from fitz import open
from PIL import Image
import pytesseract
from googletrans import Translator  # 使用 Google Translate API
import io
from docx import Document


def extract_text_and_images(pdf_path, output_docx, target_lang="zh-cn"):
	# 初始化翻译器
	translator = Translator()
	# 打开 PDF 文件
	doc = open(pdf_path)
	# 创建一个 Word 文档
	docx_file = Document()
	docx_file.add_heading('翻译结果', level=1)

	for page_num in range(len(doc)):
		page = doc.load_page(page_num)  # 加载页面
		text = page.get_text()  # 提取页面文字
		print(f"正在处理第 {page_num + 1} 页...")

		# 翻译页面文字
		if text is not None and text.strip():
			translated_text = translator.translate(text, dest=target_lang).text
			docx_file.add_heading(f'第 {page_num + 1} 页 - 文本内容', level=2)
			docx_file.add_paragraph(translated_text)

		# 提取页面中的图片

		images = page.get_images(full=True)
		if images:
			for img_index, img in enumerate(images):
				xref = img[0]
				base_image = doc.extract_image(xref)
				image_bytes = base_image["image"]
				width = base_image["width"]
				height = base_image["height"]

				# 设置一个阈值来选择较大的图像（例如，宽度大于500px）
				if width > 200 and height > 200:
					image = Image.open(io.BytesIO(image_bytes))
				else:
					continue
				# 使用 OCR 对图片文字进行识别
				ocr_text = pytesseract.image_to_string(image, lang="eng")  # 英语 OCR
				if ocr_text is not None and ocr_text.strip():
					translated_ocr_text = translator.translate(ocr_text, dest=target_lang).text

					# 保存图片到 Word
					docx_file.add_heading(f'第 {page_num + 1} 页 - 图片 {img_index + 1}', level=2)
					image.save(f"page_{page_num + 1}_image_{img_index + 1}.png")  # 可选：保存图片
					docx_file.add_paragraph("OCR 翻译内容：")
					docx_file.add_paragraph(translated_ocr_text)
	# 保存翻译后的 Word 文档
	docx_file.save(output_docx)
	print(f"翻译完成，结果已保存到 {output_docx}")

if __name__ == "__main__":
# 示例调用
	extract_text_and_images("/Users/zhanglei/xmind2testcase/docs/testlink-official-file-formats.pdf", "translated.docx", target_lang="zh-cn")