# -*- coding: utf-8 -*-
# @Time : 2025/01/09 9:17
# @Author : zhanglei
# @Email : zhspark@gmail.com
## 功能： 翻译pdf中英文，生成docx


import re
from PyPDF2 import PdfReader
from googletrans import Translator
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def translate_pdf_to_docx_with_font(input_pdf_path, output_docx_path, font_name="SimSun", font_size=12):
    try:
        # 初始化 PDF 阅读器、翻译器和 Word 文档
        reader = PdfReader(input_pdf_path)
        translator = Translator()
        doc = Document()

        # 设置默认样式
        style = doc.styles['Normal']
        style.font.name = font_name  # 设置字体名称
        style.font.size = Pt(font_size)  # 设置字体大小
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 设置中文字体支持

        # 循环读取 PDF 的每一页
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            text = re.sub(r"\t", " ", text)
            if text.strip():
                # 翻译英文文本为中文
                translated_text = translator.translate(text, src='en', dest='zh-cn').text

                # 将翻译的内容写入 Word 文档
                doc.add_heading(f'Page {page_num}', level=1)
                paragraph = doc.add_paragraph(translated_text)

                # 设置段落样式
                paragraph.style.font.name = font_name
                paragraph.style.font.size = Pt(font_size)
                paragraph.style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        # 保存翻译结果到新的 Word 文档
        doc.save(output_docx_path)
        print(f"Translation completed! New Word document saved to {output_docx_path}")
    except Exception as e:
        print(f"Error occurred: {e}")

if __name__ == "__main__":
    # 示例用法
    input_pdf_path = "/Users/zhanglei/Downloads/Value-of-Professional-Funds-Management.pdf"  # 输入 PDF 文件路径
    output_docx_path = "translated_output.docx"  # 输出翻译后的 Word 文件路径
    translate_pdf_to_docx_with_font(input_pdf_path, output_docx_path, font_name="SimSun", font_size=12)