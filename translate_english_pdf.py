# -*- coding: utf-8 -*-
# @Time : 2025/01/09 9:17
# @Author : zhanglei
# @Email : zhspark@gmail.com
## 功能： 翻译英文pdf成指定语言，生成指定文档
import re
from PyPDF2 import PdfReader
from googletrans import Translator
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.pagesizes import letter

def translate_pdf_to_docx_with_font(input_pdf_path, output_docx_path, font_name="SimSun", font_size=12, language="zh-cn"):
    try:
        # 初始化 PDF 阅读器、翻译器和 Word 文档
        reader = PdfReader(input_pdf_path)
        translator = Translator()
        doc = Document()

        # 设置默认样式
        style = doc.styles['Normal']
        style.font.name = font_name  # 设置字体名称
        style.font.size = Pt(font_size)  # 设置字体大小
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 设置中文字体支持

        # 循环读取 PDF 的每一页
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            text = re.sub(r"\t", " ", text)
            if text.strip():
                # 翻译英文文本为中文
                translated_text = translator.translate(text, src='en', dest=language).text

                # 将翻译的内容写入 Word 文档
                doc.add_heading(f'Page {page_num}', level=1)
                paragraph = doc.add_paragraph(translated_text)

                # 设置段落样式
                paragraph.style.font.name = font_name
                paragraph.style.font.size = Pt(font_size)
                paragraph.style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        # 保存翻译结果到新的 Word 文档
        doc.save(output_docx_path)
        print(f"Translation completed! New Word document saved to {output_docx_path}")
    except Exception as e:
        print(f"Error occurred: {e}")

def translate_pdf_to_new_pdf(input_pdf_path, output_pdf_path, font_path="SimSun.ttf", language="zh-cn"):
    try:
        # 注册中文字体
        pdfmetrics.registerFont(TTFont('SimSun', font_path))

        # 初始化 PDF 阅读器和翻译器
        reader = PdfReader(input_pdf_path)
        translator = Translator()

        # 创建新的 PDF 文件
        c = canvas.Canvas(output_pdf_path, pagesize=letter)
        width, height = letter
        margin = 40

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            text = re.sub(r"\t", " ", text)
            if text.strip():
                # 翻译文本
                translated_text = translator.translate(text, src='en', dest=language).text

                # 使用注册的中文字体
                c.setFont('SimSun', 12)

                # 分段处理翻译内容，避免超出页面宽度
                lines = translated_text.split("\n")
                y = height - margin
                for line in lines:
                    wrapped_lines = [line[i:i+40] for i in range(0, len(line), 40)]
                    for wrapped_line in wrapped_lines:
                        if y < margin:
                            c.showPage()  # 创建新页面
                            c.setFont('SimSun', 12)
                            y = height - margin
                        c.drawString(margin, y, wrapped_line)
                        y -= 20

                # 创建新页面
                c.showPage()
        c.save()
        print(f"Translation completed! New PDF saved to {output_pdf_path}")
    except Exception as e:
        print(f"Error occurred: {e}")

def translate_pdf_to_txt(pdf_path, output_path, language="zh-cn"):
    # Step 1: Read the PDF file
    try:
        reader = PdfReader(pdf_path)
        translator = Translator()
        translated_content = ""

        # Step 2: Extract text from each page and translate
        page_num = 0
        for page in reader.pages:
            page_num += 1
            text = page.extract_text()
            if text is not None and text.strip():
                # Translate English to Chinese
                text = re.sub(r"\t", " ", text)
                trans = translator.translate(str(text), src='en', dest=language)
                translated_text = trans.text
                translated_content += translated_text + "\n------Page: {}--------\n".format(page_num)

        print(translated_content)
        # Step 3: Save the translated content to a new file
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(translated_content)
        print(f"Translation completed! Translated content saved to {output_path}")
    except Exception as e:
        print(f"Error occurred: {e}")

def generate_file(output_path, extension_name, language):
    if extension_name == 'docx':
        output_docx_path = output_path + '.docx'
        translate_pdf_to_docx_with_font(input_pdf_path, output_docx_path, font_name="SimSun", font_size=12, language=language)
    if extension_name == 'pdf':
        output_pdf_path = output_path + '.pdf'
        translate_pdf_to_new_pdf(input_pdf_path, output_pdf_path, font_path="SimSun.ttf", language=language)
    if extension_name == 'txt':
        output_txt_path = output_path + '.txt'
        translate_pdf_to_txt(input_pdf_path, output_txt_path, language=language)



if __name__ == "__main__":
    # 示例用法
    input_pdf_path = "/Users/zhanglei/Downloads/Value-of-Professional-Funds-Management.pdf"  # 输入 PDF 文件路径
    output_path = "translated_output"  # 输出翻译后的 Word 文件路径
    ## 用法：
        # 翻译成中文docx
        ## generate_file(output_path, extension_name="docx", language="zh-cn")
        # 翻译成日文pdf
        ## generate_file(output_path, extension_name="pdf", language="ja")
        # 翻译成中文txt
        ## generate_file(output_path, extension_name="txt", language="zh-cn")
    generate_file(output_path, extension_name="txt", language="ja")
